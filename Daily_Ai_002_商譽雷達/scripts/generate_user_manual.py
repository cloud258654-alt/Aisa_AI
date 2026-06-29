"""
Generate Sentinel ECXIP Software User Manual (.docx)
Detailed operation guide for end users.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for lv in range(1, 5):
    h = doc.styles[f'Heading {lv}']
    h.font.name = 'Calibri'; h.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
    sizes = {1:22, 2:16, 3:13, 4:11.5}
    h.font.size = Pt(sizes[lv])

def T(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="0071E3"/>'))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
            if ri%2==0: c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F7"/>'))
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def S(n, title, desc, subs=None):
    p = doc.add_paragraph()
    r = p.add_run(f'步驟 {n}：{title}')
    r.font.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x00,0x71,0xE3)
    doc.add_paragraph(desc)
    if subs:
        for s in subs:
            bp = doc.add_paragraph(s, style='List Bullet')
            for r in bp.runs: r.font.size = Pt(10.5)

def N(text, t='info'):
    colors = {'info':('E8F0FE','1D1D1F'),'warning':('FFF3E0','E65100'),'tip':('E8F5E9','2E7D32')}
    bg,fg = colors.get(t, colors['info'])
    p = doc.add_paragraph()
    r = p.add_run(f'  {text}')
    r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = RGBColor(*[int(fg[i:i+2],16) for i in range(0,6,2)])
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)

def IMG(cap):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[ 圖：{cap} ]')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x86,0x86,0x8B); r.font.italic = True
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)

# ══════ COVER ══════
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sentinel AI ECXIP'); r.font.size = Pt(34); r.font.bold = True; r.font.color.rgb = RGBColor(0x00,0x71,0xE3)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('軟體操作說明書\nSoftware User Manual'); r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x51,0x51,0x54)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'版本：v1.0.0 Enterprise MVP\n日期：{datetime.date.today().strftime("%Y-%m-%d")}\n適用對象：品牌經理、營運主管、客服團隊、高階決策者')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x86,0x86,0x8B)
doc.add_page_break()

# ══════ TOC ══════
doc.add_heading('目錄', level=1)
for item in ['1. 系統概述','2. 登入與啟動','3. 介面導覽','4. 品牌決策駕駛艙','5. 晨報決策中心','6. 門市戰力排行','7. 7日預測中心','8. AI 學習記憶','9. 顧客旅程診斷','10. 輿情聲量即時串流','11. AI 品牌經理指揮中心','12. AI 語意分析沙盒','13. 語言切換','14. 系統健康監控','15. 常見操作流程','16. 疑難排解']:
    p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.size = Pt(10)
doc.add_page_break()

# ══════ 1 ══════
doc.add_heading('1. 系統概述', level=1)
doc.add_paragraph('Sentinel AI ECXIP (Enterprise Customer Experience Intelligence Platform) 是一套企業級顧客體驗智慧平台，協助品牌管理者即時掌握全品牌健康狀態、監控社群輿情、診斷顧客旅程摩擦點，並透過 AI 獲得改善建議。')
doc.add_paragraph('本手冊適用於 Sentinel ECXIP v1.0.0 Enterprise MVP 版本，涵蓋所有已上線功能模組的完整操作說明。')
doc.add_heading('系統需求', level=2)
T(['項目','需求'],[['瀏覽器','Google Chrome 90+ / Firefox 90+ / Microsoft Edge 90+'],['螢幕解析度','建議 1920×1080 以上，最低 1366×768'],['網路連線','需連接至系統伺服器 (預設 http://localhost:26117)'],['帳號權限','需具備有效使用者帳號與密碼']], widths=[3,13])
doc.add_page_break()

# ══════ 2 ══════
doc.add_heading('2. 登入與啟動', level=1)
doc.add_heading('2.1 Docker 啟動 (系統管理員)', level=2)
S(1,'確認 Docker Desktop 已安裝並執行','請確認您的電腦已安裝 Docker Desktop，且 Docker 引擎正在執行中。')
S(2,'執行 Docker Compose','開啟終端機，切換至專案的 docker 目錄，執行啟動指令：',['cd docker','docker compose up -d'])
S(3,'確認服務狀態','執行 docker compose ps 確認所有 7 個服務 (postgres / redis / backend / celery_worker / celery_beat / frontend / nginx) 狀態均為 Up。')
N('首次啟動可能需要數分鐘下載映像檔，請耐心等待。')
doc.add_heading('2.2 本地開發啟動', level=2)
S(1,'啟動後端','cd backend && pip install -r requirements.txt && python main.py')
S(2,'啟動前端','另開終端機：cd frontend && python -m http.server 26117')
S(3,'產生展示資料 (可選)','cd backend && python scripts/seed_demo_data.py')
doc.add_heading('2.3 開啟系統', level=2)
S(1,'開啟瀏覽器','在瀏覽器網址列輸入 http://localhost:26117')
S(2,'系統載入','畫面將顯示 Sentinel ECXIP 的 Apple Frosted Glass 風格主介面。若頂部顯示橘色「DEMO MODE」橫幅，表示目前為示範模式。')
N('Demo 測試帳號：admin@sentinel.ai / demo123','tip')
IMG('Sentinel ECXIP 系統首頁')
doc.add_page_break()

# ══════ 3 ══════
doc.add_heading('3. 介面導覽', level=1)
doc.add_paragraph('系統介面由三個主要區域組成：左側導覽列、頂部操作列、中央工作區。')
doc.add_heading('3.1 左側導覽列', level=2)
T(['導覽項目','功能說明'],[['品牌決策總覽','顯示五大核心品牌指標與風險指示條'],['晨報決策中心','AI COO 每日晨報、今日警示、門市排行摘要'],['門市戰力排行','全品牌門市健康排行、危險/改善/惡化分類'],['7日預測中心','品牌健康/風險/負評量 7 日預測與營運模擬'],['AI 學習記憶','歷史案例搜尋、模式發現、新案例提交'],['顧客旅程診斷','六個接觸點滿意度視覺化與摩擦診斷'],['聲量即時串流','6 大渠道即時評論串流與輿情處理'],['AI 分析沙盒','自定義文字 AI 語意分析與改善建議']], widths=[3.5,12.5])
doc.add_paragraph('點擊任一導覽項目，工作區將自動捲動至對應頁面區塊。當前頁面項目會以藍色高亮顯示。導覽列底部顯示 AI Engine Live 連線狀態及系統健康摘要。')
doc.add_heading('3.2 頂部操作列', level=2)
T(['元件','功能'],[['搜尋框','關鍵字搜尋門市、事件、輿情（目前為預留功能）'],['門市篩選器','下拉選單切換監控門市：全品牌 / 信義旗艦店 / 忠孝 SOGO / 板橋大遠百 / 台中公益店'],['系統時間','即時顯示目前日期與時間'],['使用者頭像','顯示目前登入角色 (Brand Manager)'],['語言切換器','繁中 / EN 按鈕，點擊即時切換介面語言']], widths=[3,13])
doc.add_heading('3.3 中央工作區', level=2)
doc.add_paragraph('中央工作區為可捲動內容區域，由上至下依序排列各功能模組：指標儀表板 → 晨報決策中心 → 門市排行 → 預測中心 → 學習記憶 → 即時串流 (左) + 旅程診斷 (右) → AI 指揮中心 → NLP 沙盒')
IMG('系統介面全覽圖')
doc.add_page_break()

# ══════ 4 ══════
doc.add_heading('4. 品牌決策駕駛艙 Brand Cockpit', level=1)
doc.add_paragraph('系統首頁頂部的五張指標卡片，提供即時的全品牌健康狀態概覽。每張卡片包含指標名稱、數值、變化趨勢與圓環進度圖。')
IMG('品牌決策駕駛艙 — 五大指標卡片')
doc.add_heading('4.1 指標說明', level=2)
T(['指標','顯示格式','說明'],[['品牌健康度','92.4 (數字)','綜合 AI 分析結果的加權分數 (0-100)，含上週比較'],['門市健康指數','98.1%','所有門市的加權平均健康分數，含昨日比較'],['顧客滿意度','4.82 / 5','跨渠道 (Google 為主) 加權平均星級評分，附星號圖示'],['危機解決率','94.5%','品牌危機事件的結案比例，含上月比較'],['商譽風險值','Low (12)','即時風險分數 0-100，長條顏色自動變換']], widths=[3,2.5,10.5])
doc.add_heading('4.2 風險指示條', level=2)
T(['風險等級','分數範圍','顏色'],[['低風險','0 - 29','綠色'],['中風險','30 - 69','橘色'],['極度危機','70 - 100','紅色']], widths=[3,3,10])
doc.add_heading('4.3 動態指標變化', level=2)
doc.add_paragraph('系統指標並非靜態數字。當即時串流出現新負評、AI 沙盒分析高風險文字、或處理輿情事件時，各項指標會自動浮動更新，提供真實的數據回饋感。')
N('Demo 模式下，負評串流每 6 秒降低品牌健康度 0.2、增加風險值 2；正面評論則反向微幅提升。','info')
doc.add_page_break()

# ══════ 5 ══════
doc.add_heading('5. 晨報決策中心 Executive Briefing Center', level=1)
doc.add_paragraph('每日自動生成的高階主管晨報，位於指標儀表板下方。無需任何操作即可了解當日品牌全貌。')
IMG('晨報決策中心完整畫面')
doc.add_heading('5.1 晨報內容區塊', level=2)
S(1,'閱讀問候與日期','頁面頂部顯示問候語（如「早安，Brand Manager」）、今日日期、風險等級徽章。')
S(2,'查看關鍵指標摘要','四個核心數字：品牌健康度、門市健康指數、商譽風險值、今日 VOC 聲量總數。')
S(3,'閱讀今日最大問題','左側問題卡片顯示今日最嚴重的品牌問題，包含描述、嚴重程度、受影響門市、AI 建議。')
S(4,'查看 AI COO 建議','顯示 AI 營運長提出的策略建議，每項附帶預期成果與信心度指標。')
S(5,'查看今日警示','右側警示面板列出今日所有風險通知，依嚴重程度分類。')
S(6,'查看門市排行摘要','Top 5 門市排行榜：含排名、名稱、健康分數、狀態、趨勢箭頭。')
S(7,'查看行動事項','今日需執行的具體行動清單。')
S(8,'查看 7 日風險預測','底部走勢圖顯示未來 7 天品牌健康度、風險分數、負評量的預測趨勢。')
doc.add_heading('5.2 重新整理', level=2)
doc.add_paragraph('點擊右上角重新整理按鈕 (圓形箭頭圖示) 可手動重新載入晨報資料。')
doc.add_page_break()

# ══════ 6 ══════
doc.add_heading('6. 門市戰力排行 Store Ranking', level=1)
doc.add_paragraph('全品牌門市的健康排行總表，協助管理者快速識別需要關注的門市。')
IMG('門市戰力排行 — 表格與分類頁籤')
doc.add_heading('6.1 查看排行', level=2)
S(1,'進入排行頁面','點擊左側導覽列「門市戰力排行」，或向下捲動至排行區塊。')
S(2,'切換分類頁籤','點擊表格上方四個頁籤按鈕：All Stores（所有門市）、Critical（危險）、Improving（改善中）、Declining（惡化中）。')
S(3,'閱讀排行表格','表格欄位：Rank（排名）、Store Name（門市名稱）、Health Score（健康分數 + 顏色條）、Risk Level（風險標籤）、Trend（趨勢箭頭 ↑→↓）、Issues（問題數）。')
doc.add_heading('6.2 查看門市詳情', level=2)
S(1,'點擊門市資料列','點擊任一門市資料列，表格下方展開該門市詳細診斷面板。')
S(2,'閱讀門市詳細診斷','詳細面板顯示各項分數、趨勢分析、問題列表與 AI 建議。')
doc.add_page_break()

# ══════ 7 ══════
doc.add_heading('7. 7日預測中心 Predictive Intelligence', level=1)
doc.add_paragraph('基於歷史數據的 AI 預測引擎，提供未來 7 天的關鍵指標趨勢預測。')
IMG('7日預測中心 — 四個預測圖表面板')
doc.add_heading('7.1 查看預測圖表', level=2)
S(1,'進入預測頁面','點擊左側導覽列「7日預測中心」。')
S(2,'閱讀四個預測面板','系統自動顯示：品牌健康度 7 日預測、商譽風險 7 日預測、負面聲量 7 日預測、預測信心度指標。')
doc.add_heading('7.2 營運模擬 (What-if)', level=2)
S(1,'輸入模擬情境','在預測圖表下方的「營運模擬」區域輸入假設情境，例如：增加 10% 尖峰時段人力配置。')
S(2,'執行模擬','點擊「執行模擬」按鈕。')
S(3,'查看模擬結果','系統顯示原始數值、模擬後數值、影響評估與信心度。')
N('預測基於加權移動平均與趨勢外推演算法。未來版本將整合機器學習模型以提高預測精準度。','info')
doc.add_page_break()

# ══════ 8 ══════
doc.add_heading('8. AI 學習記憶 Learning Memory Engine', level=1)
doc.add_paragraph('企業專屬的知識圖譜系統，記錄歷史品牌事件與解決方案，讓 AI 從過去經驗中持續學習。')
IMG('AI 學習記憶 — 相似案例與模式面板')
doc.add_heading('8.1 查看歷史案例', level=2)
S(1,'進入學習記憶頁面','點擊左側導覽列「AI 學習記憶」。')
S(2,'查看相似案例','左側面板顯示歷史案例。每張卡片含：案例標題、事件描述、解決方案、結果與成效評分。')
S(3,'查看成功模式','右上方面板顯示 AI 發現的成功模式：名稱、頻率、成功率、推薦行動。')
S(4,'查看改善建議','右下方面板顯示基於歷史數據的學習建議。')
doc.add_heading('8.2 提交新案例', level=2)
S(1,'點擊「新增案例」','點擊頁面右上角「Store New Case」按鈕，展開案例提交表單。')
S(2,'填寫案例資訊','填寫：Case Title（案例標題）、Affected Store（受影響門市，選填）、Description（事件描述、解決方式、結果）。')
S(3,'提交案例','點擊「Submit to Learning Engine」按鈕，系統將儲存案例並自動進行關鍵字提取與相似度比對。')
S(4,'取消提交','若不提交，點擊「Cancel」按鈕關閉表單。')
N('提交的案例將被 AI 自動分析關鍵字。成功案例的解決方案將影響未來的 AI 推薦。','tip')
doc.add_page_break()

# ══════ 9 ══════
doc.add_heading('9. 顧客旅程診斷 Customer Journey Map', level=1)
doc.add_paragraph('視覺化顧客完整旅程的六個關鍵接觸點，即時顯示各節點滿意度與摩擦警示。')
IMG('顧客旅程診斷 — 六節點流程圖與診斷卡片')
doc.add_heading('9.1 閱讀旅程圖', level=2)
S(1,'查看六個旅程節點','旅程圖由左至右顯示：搜尋探索 → 線上預約 → 候位到店 → 服務體驗 → 付款收銀 → 評論回饋。')
S(2,'辨識節點狀態','顏色辨識：綠色邊框 = 良好 (>90%)、橘色邊框 + 閃爍 = 警告 (70-90%)、紅色邊框 + 快速閃爍 = 嚴重 (<70%)。')
S(3,'查看節點分數','每個節點下方顯示該接觸點的滿意度百分比。')
doc.add_heading('9.2 閱讀診斷卡片', level=2)
S(1,'查看診斷卡片','旅程圖下方列出有問題的節點診斷卡片，包含：問題標題、受影響門市標籤、詳細描述。')
doc.add_heading('9.3 切換監控門市', level=2)
S(1,'使用門市篩選器','於頂部操作列的門市下拉選單中，選擇欲查看的門市。')
S(2,'觀察變化','切換門市後，旅程節點狀態、分數與診斷卡片內容會即時更新。')
N('不同門市的旅程診斷數據各自獨立，可逐一切換查看各分店的營運狀態。','info')
doc.add_page_break()

# ══════ 10 ══════
doc.add_heading('10. 輿情聲量即時串流 Voice of Customer Stream', level=1)
doc.add_paragraph('左側面板的即時評論串流，模擬來自 6 大社群渠道的顧客回饋，每則評論皆經 AI 自動標記。')
IMG('即時串流 — 評論卡片與 AI 分析標籤')
doc.add_heading('10.1 閱讀評論卡片', level=2)
S(1,'辨識來源渠道','每張卡片左上角顯示渠道圖示與顏色：Google Maps (橘黃)、Threads (黑)、Facebook (藍)、Instagram (粉紅)、Dcard (深藍)、PTT (紅)。')
S(2,'閱讀評論內容與標籤','卡片中央顯示全文，下方標籤列顯示：情感標籤 (POSITIVE 綠 / NEUTRAL 橘 / NEGATIVE 紅)、主題標籤 (藍)、門市標籤 (灰)、高風險警示 (紅邊框)。')
S(3,'查看時間戳記','卡片右上角顯示「剛剛」。')
doc.add_heading('10.2 處理輿情', level=2)
S(1,'點擊「處理輿情」按鈕','在評論卡片底部，點擊藍色「處理輿情」按鈕。')
S(2,'選擇處理方式','彈窗提供三個選項：「指派改善 SOP 給門市」、「使用 AI 產出回覆並張貼」、「向上回報為品牌危機事件」。')
S(3,'查看處理結果','選擇後系統即時更新品牌指標（解決率上升、風險下降），並在 AI 終端機記錄日誌。')
S(4,'關閉彈窗','點擊右上角「×」或點擊彈窗外灰色區域關閉。')
doc.add_heading('10.3 串流控制', level=2)
S(1,'暫停/啟動串流','點擊面板右上角暫停按鈕 (兩豎線圖示) 可暫停；再次點擊 (播放圖示) 恢復。')
S(2,'清除視窗','點擊垃圾桶圖示按鈕清空目前所有評論卡片。新評論仍會持續流入。')
doc.add_page_break()

# ══════ 11 ══════
doc.add_heading('11. AI 品牌經理指揮中心 AI Brand Manager', level=1)
doc.add_paragraph('模擬品牌危機事件，展示 AI 如何自動執行根因分析、產生 SOP、撰寫公關聲明與法務建議。')
IMG('AI 品牌經理 — 終端機與四組決策分頁')
doc.add_heading('11.1 觸發危機模擬', level=2)
S(1,'選擇危機情境','點擊三個預設事件按鈕：Threads 食安疑慮謠言、Google 負評爆發 (排隊糾紛)、PTT 員工服務態度爆料。')
S(2,'觀察 AI 分析過程','點擊後終端機以擬真打字延遲顯示 AI 分析日誌：鎖定社群貼文來源 → 回溯內部 POS 與監視器數據 → 萃取輿情情緒指標 → 觸發商譽危機預警 → 生成 AI 決策支援模組。')
S(3,'查看儀表板變化','危機觸發後，頂部指標儀表板的品牌健康度下降、風險值上升、風險長條顏色可能變為橘色或紅色。')
doc.add_heading('11.2 使用決策支援分頁', level=2)
doc.add_paragraph('終端機下方顯示四組決策分頁：')
S(1,'Root Cause (原因解析)','顯示 AI 交叉分析社群數據與內部營運數據後的根本原因診斷。包含：輿情源頭、內部數據比對結果、診斷結論。')
S(2,'Operational SOP (改善流程)','顯示互動式待辦清單。勾選每個核取方塊項目，當全部完成後系統自動提升品牌健康度並降低風險值。',['每勾選一個項目，終端機顯示進度更新訊息','全部完成後顯示「所有改善 SOP 已全數確認，品牌健康度回升」'])
S(3,'PR Statement (公關回應)','顯示 AI 自動生成公關聲明。操作：點擊「複製聲明」複製至剪貼簿；點擊「一鍵張貼至...」模擬發布至社群平台。')
S(4,'Legal & Training (法務與培訓)','包含：法務防線警示（法律風險提醒）與門市再培訓要點（員工訓練重點）。')
N('每次僅能觸發一個危機情境。觸發新危機時，前次模擬結果將被清除。','warning')
doc.add_page_break()

# ══════ 12 ══════
doc.add_heading('12. AI 語意分析沙盒 NLP Sandbox', level=1)
doc.add_paragraph('手動輸入任意文字，即時體驗 Sentinel 的多維度語意分析能力。')
IMG('NLP 沙盒 — 分析結果與改善建議')
doc.add_heading('12.1 執行文字分析', level=2)
S(1,'輸入文字','在文字輸入框中輸入欲分析的顧客評論或社群貼文。可點擊「填入示範負評」自動填入預設示範文字。')
S(2,'執行分析','點擊藍色「執行 AI 深度分析」按鈕。')
S(3,'等待掃描動畫','顯示光線掃描動畫與「Sentinel NLP Pipeline Processing...」提示，約 1.5 秒。')
S(4,'查看分析結果','掃描完成後顯示四個結果方塊與改善建議區塊。')
doc.add_heading('12.2 閱讀分析結果', level=2)
T(['分析項目','顯示內容'],[['情感分析 Sentiment','情感傾向 + 信心度 (如 Negative 96%)'],['細粒度情緒 Emotion','具體情緒類別 + 表情符號 (如 Anger / Joy)'],['旅程痛點 Touchpoint','對應的顧客旅程接觸點 (如 Wait & Service)'],['商譽風險等級 Risk','風險等級 + 0-100 分數 (如 High Risk 72)']], widths=[5,11])
doc.add_heading('12.3 閱讀改善建議', level=2)
S(1,'復原回覆建議 (PR Draft)','AI 針對輸入文字自動生成的公關回覆草稿，可直接複製使用。')
S(2,'營運改善 SOP 建議','AI 針對識別的問題自動生成具體的門市營運改善步驟。')
N('沙盒分析使用關鍵字規則引擎。包含「爛/冰/臭臉/等很久/拉肚子/差」等關鍵字將判定為負面情感。','info')
N('沙盒分析結果會即時影響頂部儀表板的品牌指標。高風險分析將降低品牌健康度並提高風險值。','warning')
doc.add_page_break()

# ══════ 13 ══════
doc.add_heading('13. 語言切換 Language Switcher', level=1)
doc.add_paragraph('系統支援繁體中文 (zh-TW) 與英文 (en-US) 雙語介面，可即時切換。')
S(1,'找到語言切換器','語言切換器位於頂部操作列最右側，顯示為「繁中」或「EN」文字方塊。')
S(2,'切換語言','點擊「繁中」或「EN」按鈕即可在兩種語言之間切換。')
S(3,'觀察變化','切換後所有介面文字立即更新，無需刷新頁面。')
S(4,'確認語言保存','關閉瀏覽器後重新開啟系統，語言設定自動沿用上次選擇。')
N('若找不到對應翻譯，系統將依序回退至英文 → Key 名稱，不會造成畫面錯誤。','info')
doc.add_page_break()

# ══════ 14 ══════
doc.add_heading('14. 系統健康監控 System Health', level=1)
doc.add_paragraph('側邊欄底部提供即時的系統服務健康狀態摘要。')
S(1,'找到健康監控面板','系統健康面板位於左側導覽列最底部，AI Engine Live 連線狀態下方。')
S(2,'閱讀健康狀態','面板顯示五項核心服務狀態，以顏色圓點標示：System (API)、DB (PostgreSQL)、Redis、AI (Router)、Queue (Celery)。')
S(3,'辨識狀態顏色','綠色 = Healthy（正常）、橘色 = Warning（異常但可用）、紅色 = Error（中斷）。')
doc.add_heading('14.1 API 健康檢查', level=2)
T(['API 端點','說明'],[['GET /api/v1/system/health','全系統深度檢查 (DB/Redis/Celery/WebSocket/AI/Crawler/Notification)'],['GET /api/v1/system/status','快速存活檢查'],['GET /api/v1/health','基礎健康檢查']], widths=[6,10])
doc.add_page_break()

# ══════ 15 ══════
doc.add_heading('15. 常見操作流程', level=1)
doc.add_heading('15.1 每日品牌巡檢流程', level=2)
S(1,'查看晨報','進入晨報決策中心，閱讀 AI COO 摘要、今日問題與警示。')
S(2,'檢查指標','查看品牌駕駛艙五大指標，確認風險值是否正常。')
S(3,'查看排行','進入門市排行，找出危險或惡化中的門市。')
S(4,'深入診斷','切換門市篩選器至問題門市，查看顧客旅程診斷。')
S(5,'處理輿情','在即時串流中找到相關負評，點擊處理。')
S(6,'查看預測','確認 7 日預測趨勢，評估是否需要預防性措施。')
S(7,'查看學習','搜尋歷史相似案例，參考過去的成功解決方案。')

doc.add_heading('15.2 危機應變流程', level=2)
S(1,'觸發模擬','在 AI 品牌經理點擊對應的危機情境按鈕。')
S(2,'閱讀根因','切換至 Root Cause 分頁，了解事件根本原因。')
S(3,'執行 SOP','切換至 Operational SOP 分頁，逐一勾選改善項目。')
S(4,'發布公關','切換至 PR Statement 分頁，複製或模擬發布公關聲明。')
S(5,'確認法務','切換至 Legal & Training 分頁，確認法務提醒與培訓要點。')
S(6,'確認修復','當所有 SOP 項目勾選完成後，確認品牌指標已回升。')

doc.add_heading('15.3 AI 分析驗證流程', level=2)
S(1,'輸入測試文字','在 NLP 沙盒中輸入欲測試的評論文字。')
S(2,'執行分析','點擊執行分析，查看情感/情緒/痛點/風險結果。')
S(3,'查看建議','閱讀 AI 生成的 PR 回覆草案與 SOP 改善建議。')
S(4,'記錄案例','若為真實案例，進入學習記憶提交至學習引擎。')
doc.add_page_break()

# ══════ 16 ══════
doc.add_heading('16. 疑難排解', level=1)
T(['問題','可能原因','解決方法'],[
['頁面無法開啟','Docker 服務未啟動','確認 docker compose ps 所有服務 Up，或重新執行 docker compose up -d'],
['儀表板數據不更新','串流已暫停','點擊串流暫停按鈕恢復串流'],
['語言切換無效','i18n 腳本載入失敗','重整頁面 (F5)，確認瀏覽器主控台無錯誤'],
['側邊欄顯示紅色連線','後端 API 服務中斷','檢查後端容器狀態：docker compose logs backend'],
['晨報無內容','後端服務未返回資料','確認後端 API 正常：curl http://localhost:8000/api/v1/health'],
['門市排行空白','尚無排行資料','點擊重新整理或等待排程任務自動計算（每日 3:00 AM）'],
['學習記憶無歷史案例','尚未提交學習案例','點擊新增案例按鈕，建立第一筆學習案例'],
['預測圖表無資料','歷史數據不足','至少需要 7 天以上的歷史數據才能生成預測'],
['沙盒分析無反應','未輸入文字','請先輸入至少 5 個字以上的評論文字'],
['Demo 橫幅不正確','Demo Mode 設定問題','確認 .env 中 DEMO_MODE=true'],
], widths=[4.5,4.5,7])

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Sentinel AI ECXIP v1.0.0 軟體操作說明書 —')
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x00,0x71,0xE3)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('文件版本 1.0 | 適用系統版本 v1.0.0 Enterprise MVP | Feature Freeze')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x86,0x86,0x8B)

# ── Save ──
out = r'D:\Ai study\Aisa_AI\Daily_Ai_002_商譽雷達\Sentinel_ECXIP_操作說明書_v1.0.0.docx'
doc.save(out)
print(f'Saved: {out}')
