"""
Generate Sentinel ECXIP Project Proposal Document (.docx)
Client-facing feature introduction and technical overview.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_colored_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0071E3"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
            if r % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F7"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('Sentinel AI ECXIP')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x71, 0xE3)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run(
    'Enterprise Customer Experience Intelligence Platform\n'
    '企業級顧客體驗智慧平台'
)
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x51, 0x51, 0x54)

doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run(
    f'專案計劃書 / Project Proposal\n\n'
    f'版本：v1.0.0 Enterprise MVP\n'
    f'日期：{datetime.date.today().strftime("%Y-%m-%d")}\n'
    f'狀態：Release Candidate — Ready for Demo'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
doc.add_heading('目錄 Table of Contents', level=1)
toc_items = [
    '1. 專案摘要 Executive Summary',
    '2. 產品定位 Product Positioning',
    '3. 系統架構 System Architecture',
    '4. 功能模組介紹 Feature Modules',
    '   4.1 品牌決策駕駛艙 Brand Cockpit',
    '   4.2 晨報決策中心 Executive Briefing Center',
    '   4.3 顧客之聲即時串流 Voice of Customer Stream',
    '   4.4 顧客旅程診斷 Customer Journey Map',
    '   4.5 AI 品牌經理決策指揮中心 AI Brand Manager',
    '   4.6 AI 語意分析沙盒 NLP Sandbox',
    '   4.7 門市戰力排行 Store Intelligence Ranking',
    '   4.8 7 日預測中心 Predictive Intelligence',
    '   4.9 AI 學習記憶 Learning Memory Engine',
    '   4.10 營運智慧 Operational Intelligence',
    '5. AI 人工智慧架構 AI Agent Architecture',
    '6. 技術架構 Technology Stack',
    '7. 部署方式 Deployment',
    '8. 多語系支援 i18n Internationalization',
    '9. 系統健康監控 System Health Monitoring',
    '10. 未來路線圖 Roadmap',
    '11. 結語 Closing Statement',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. 專案摘要 Executive Summary', level=1)

doc.add_paragraph(
    'Sentinel AI ECXIP (Enterprise Customer Experience Intelligence Platform) '
    '是一套專為連鎖品牌與零售企業打造的企業級顧客體驗智慧平台。'
)
doc.add_paragraph(
    '本平台整合 AI Agent、顧客之聲 (Voice of Customer)、顧客體驗 (Customer Experience)、'
    '品牌智慧 (Brand Intelligence) 與營運智慧 (Operational Intelligence)，'
    '協助企業從大量顧客回饋中自動發現問題、分析根本原因、提供改善建議，並持續追蹤改善成效。'
)
doc.add_paragraph(
    '不同於傳統商譽監控工具僅能通知企業「發生了什麼」，Sentinel ECXIP 更能回答：\n'
    '• 為什麼會發生？(Root Cause)\n'
    '• 哪些門市最需要改善？(Store Ranking)\n'
    '• 哪些流程造成顧客流失？(CX Journey Friction)\n'
    '• 下一步應如何改善？(AI Recommendations)\n'
    '• 是否已經改善成功？(Resolution Tracking)\n'
    '• 未來是否可能發生品牌危機？(Predictive Intelligence)'
)

# Key Metrics
doc.add_heading('平台關鍵指標', level=2)
add_colored_table(doc,
    ['指標', '數值', '說明'],
    [
        ['API 端點數', '60+', '15 個路由模組，覆蓋 VOC/CX/Brand/Workflow/Knowledge/Prediction 等'],
        ['AI Agent 數', '14', '專業化 AI 代理，由 Orchestrator 協調運作'],
        ['資料表數量', '31', '跨越 7 個業務領域 (DDD Bounded Contexts)'],
        ['前端元件數', '12+', '模組化 Vanilla JS SPA，Apple Frosted Glass 設計'],
        ['Docker 服務數', '7', 'PostgreSQL + Redis + Backend + Celery Worker/Beat + Frontend + Nginx'],
        ['翻譯 Key 數', '260+', '繁體中文 (zh-TW) + 英文 (en-US) 完整支援'],
        ['Celery 排程任務', '12', '每日品牌計算、晨報生成、風險預測、學習模式更新'],
        ['總程式檔案數', '166', 'Clean Architecture + DDD + SOLID 設計'],
    ],
    col_widths=[4, 2.5, 9.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 2. PRODUCT POSITIONING
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. 產品定位 Product Positioning', level=1)

doc.add_paragraph(
    'Sentinel ECXIP 不是一套單純的社群監控工具或 AI 自動回覆工具。'
    '我們的定位是 Enterprise Decision Intelligence Platform（企業決策智慧平台）。'
)

add_colored_table(doc,
    ['傳統商譽監控工具', 'Sentinel ECXIP'],
    [
        ['發現負評 → 通知 → 結束', 'VOC 收集 → AI 分析 → CX 診斷 → Root Cause → 風險評估 → AI 推薦 → 工作流程 → 追蹤改善'],
        ['被動回應', '主動預測與預防'],
        ['單一資料來源', '整合 6+ 渠道 VOC + POS + 排班 + 流量'],
        ['人工判斷', '14 個 AI Agent 自動分析與推薦'],
        ['無學習能力', '歷史案例記憶與模式發現'],
        ['無門市比較', '每日門市排行與趨勢追蹤'],
    ],
    col_widths=[5, 11]
)

doc.add_paragraph(
    '產品核心流程：\n'
    'Voice of Customer → Customer Experience → Brand Intelligence → '
    'AI Decision Support → Operational Intelligence → Continuous Improvement\n'
    '形成完整的企業智慧閉環。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. 系統架構 System Architecture', level=1)

doc.add_paragraph(
    'Sentinel ECXIP 採用分層式企業級架構設計，從前端到資料層共六層，'
    '符合 Clean Architecture 與 Domain-Driven Design (DDD) 原則。'
)

add_colored_table(doc,
    ['架構層', '技術元件', '說明'],
    [
        ['Gateway Layer', 'Nginx Reverse Proxy', '反向代理、靜態資源快取、WebSocket 升級'],
        ['API Layer', 'FastAPI (Python 3.12)', 'REST API + WebSocket + OpenAPI 自動文件'],
        ['Service Layer', '17 個業務服務', 'VOC/CX/Brand Health/Root Cause/RAG/Executive/Trends 等'],
        ['AI Layer', '14 AI Agents + Orchestrator + AI Router', '多層次 AI 分析管線、5-tier 模型選擇、成本優化'],
        ['Async Layer', 'Celery + Redis', '4 個專用佇列、12 個排程任務、非同步爬蟲與分析'],
        ['Data Layer', 'PostgreSQL 16 + Redis 7', '31 個資料表、7 個領域、快取與 Pub/Sub 訊息'],
    ],
    col_widths=[3, 5, 8]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 4. FEATURE MODULES
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. 功能模組介紹 Feature Modules', level=1)

# 4.1 Brand Cockpit
doc.add_heading('4.1 品牌決策駕駛艙 Brand Cockpit', level=2)
doc.add_paragraph(
    '即時顯示品牌五大核心指標，提供企業總部一目瞭然的決策資訊：'
)
add_colored_table(doc,
    ['指標名稱', '英文', '說明'],
    [
        ['品牌健康度', 'Brand Health Score', '綜合情感分析、滿意度、回覆率、正面佔比的加權分數'],
        ['門市健康指數', 'Store Health Index', '各門市評分、負評率、回覆品質的複合指標'],
        ['顧客滿意度', 'Customer CSAT', 'Google Rating 等跨渠道加權平均評分'],
        ['危機解決率', 'Crisis Resolution Rate', '品牌危機事件的結案比例'],
        ['商譽風險值', 'Reputation Risk Score', '即時風險指標 (0-100)，動態變換顏色 (綠→橘→紅)'],
    ],
    col_widths=[3, 3.5, 9.5]
)
doc.add_paragraph(
    '指標隨實時輿情事件與沙盒分析結果動態浮動，提供真實的指標回饋感。'
    '風險指示條依據危機程度自動變換顏色：綠色 (低風險 <30)、橘色 (中風險 30-70)、紅色 (極度危機 >70)。'
)

# 4.2 Executive Briefing
doc.add_heading('4.2 晨報決策中心 Executive Briefing Center', level=2)
doc.add_paragraph(
    '每日自動生成高階主管晨報，無需任何操作即可掌握當日品牌狀況。'
)
doc.add_paragraph(
    '晨報內容包含：\n'
    '• 今日品牌指標摘要 (健康度 / 風險 / VOC 聲量)\n'
    '• 今日最大問題與受影響門市\n'
    '• AI COO 策略建議與預期成果\n'
    '• Top 5 門市健康排行\n'
    '• 7 日風險預測走勢\n'
    '• 今日行動事項清單\n'
    '• 即時警示通知'
)
doc.add_paragraph(
    'AI COO Agent 綜合所有分析結果，每日產出繁體中文高階摘要，'
    '包含具體行動建議與預期業務影響。'
)

# 4.3 VOC
doc.add_heading('4.3 顧客之聲即時串流 Voice of Customer Stream', level=2)
doc.add_paragraph(
    '即時串流來自 6 大社群與評論渠道的顧客回饋：'
)
doc.add_paragraph(
    '• Google Reviews | • Threads | • Facebook\n'
    '• Instagram | • PTT | • Dcard'
)
doc.add_paragraph(
    '每則評論自動透過 AI Pipeline 進行：\n'
    '• 情感分析 (Sentiment: Positive / Neutral / Negative)\n'
    '• 細粒度情緒識別 (Emotion: 喜悅、憤怒、無奈、好奇等)\n'
    '• 旅程觸點分類 (Topic: 服務態度、出餐速度、候位時間等)\n'
    '• 所屬分店標記\n'
    '• 商譽風險權重標記 (Low / Mid / High / Critical)'
)
doc.add_paragraph(
    '支援暫停/啟動串流、清除視窗、以及點擊「處理輿情」彈窗進行 SOP 指派、AI 回覆產出與公關上報。'
)

# 4.4 CX Journey
doc.add_heading('4.4 顧客旅程診斷 Customer Journey Map', level=2)
doc.add_paragraph(
    '視覺化完整顧客旅程六個關鍵接觸點 (Touch Point)：'
)
doc.add_paragraph(
    '搜尋探索 → 線上預約 → 候位到店 → 服務體驗 → 付款收銀 → 評論回饋'
)
doc.add_paragraph(
    '當特定節點出現摩擦時，節點會顯示黃色警告或紅色警報並持續閃爍，'
    '下方列出詳細的 AI 診斷卡片，包含問題描述、受影響門市與關鍵數據。'
    '支援分店切換篩選，可獨立查看各門市的旅程診斷狀態。'
)

# 4.5 AI Brand Manager
doc.add_heading('4.5 AI 品牌經理決策指揮中心 AI Brand Manager', level=2)
doc.add_paragraph(
    '提供完整的品牌危機應變模擬與 AI 決策支援系統。'
    '內建三組預設品牌危機情境：'
)
doc.add_paragraph(
    '• Threads 食安疑慮謠言\n'
    '• Google 排隊糾紛負評爆發\n'
    '• PTT 員工服務態度爆料'
)
doc.add_paragraph(
    '點選後虛擬終端以擬真延遲打印 AI 蒐集分析日誌，並產生四組專屬決策模組：'
)
add_colored_table(doc,
    ['決策模組', '功能'],
    [
        ['Root Cause 原因解析', '交叉比對社群輿情與內部 POS/監視器數據，定位流程摩擦點'],
        ['Operational SOP 改善流程', '互動式待辦清單，全數勾選後系統自動修復品牌指標'],
        ['PR Statement 公關回應', '自動生成官方聲明，支援一鍵複製與社群平台張貼模擬'],
        ['Legal & Training 法務培訓', '提供法務避險建議與門市員工再培訓口訣'],
    ],
    col_widths=[5, 11]
)

# 4.6 NLP Sandbox
doc.add_heading('4.6 AI 語意分析沙盒 NLP Sandbox', level=2)
doc.add_paragraph(
    '互動式 NLP 遊樂場，開放輸入任何自定義評論文字，即時體驗多維度 AI 分析管道：'
)
doc.add_paragraph(
    '• 情感分析結果與信心度\n'
    '• 細粒度情緒 (喜悅 / 憤怒 / 無奈 / 好奇等)\n'
    '• 旅程痛點定位\n'
    '• 商譽風險等級 (0-100)\n'
    '• 自動生成對應的 PR 官方回覆範本\n'
    '• 自動生成門市營運 SOP 改善建議'
)
doc.add_paragraph(
    '點擊分析後播放科技感光線掃描動畫，即時展示分析結果。'
    '支援關鍵字規則引擎，可準確辨識中文負評中的多種情緒與問題類型。'
)

# 4.7 Store Intelligence
doc.add_heading('4.7 門市戰力排行 Store Intelligence Ranking', level=2)
doc.add_paragraph(
    '每日自動計算各門市的健康分數、CX 分數、風險等級，並進行全品牌排行。'
)
doc.add_paragraph(
    '功能特色：\n'
    '• 四大分類檢視：所有門市 / 危險門市 / 改善中門市 / 惡化中門市\n'
    '• 健康分數顏色條與趨勢箭頭\n'
    '• 點擊展開門市詳細診斷\n'
    '• 風險等級即時標記\n'
    '• 問題數統計'
)

# 4.8 Predictions
doc.add_heading('4.8 7 日預測中心 Predictive Intelligence', level=2)
doc.add_paragraph(
    '基於歷史數據進行 7 日多因子預測，協助企業提前準備：'
)
doc.add_paragraph(
    '• 品牌健康度 7 日預測\n'
    '• 商譽風險分數趨勢\n'
    '• 負面聲量預測\n'
    '• 預測信心度指標'
)
doc.add_paragraph(
    '支援「What-if 營運模擬」功能：輸入假設情境 (如增加人力、調整價格)，'
    '系統模擬預測對品牌指標的影響。'
)

# 4.9 Learning Memory
doc.add_heading('4.9 AI 學習記憶 Learning Memory Engine', level=2)
doc.add_paragraph(
    '建立企業專屬的知識圖譜，讓 AI 從歷史案例中持續學習，不重複犯錯：'
)
doc.add_paragraph(
    '• 歷史案例儲存與相似度搜尋 (Jaccard 關鍵字匹配)\n'
    '• AI 自動發現成功模式 (Pattern Discovery)\n'
    '• 「曾成功解決方式」推薦\n'
    '• 解決方案有效性追蹤\n'
    '• 新案例一鍵提交至學習引擎'
)
doc.add_paragraph(
    '範例：當發生「等候時間過長」事件時，AI 自動搜尋過去相似案例，'
    '發現去年透過增加晚班人力成功解決，系統推薦同樣策略。'
)

# 4.10 Operational Intelligence
doc.add_heading('4.10 營運智慧 Operational Intelligence', level=2)
doc.add_paragraph(
    '將 VOC 事件與企業營運數據連結，找出隱藏的根因關聯：'
)
doc.add_paragraph(
    '整合資料來源：POS 銷售 / 訂單量 / 門市流量 / 員工排班 / 庫存狀態 / 促銷活動 / 客訴單量 / 服務產能'
)
doc.add_paragraph(
    '分析邏輯範例：\n'
    '顧客抱怨等待太久 → 查 POS 18:00 訂單量暴增 220% → 查排班晚班只有 2 人 → '
    'AI 判斷非服務態度問題而是尖峰人力不足 → 建議 18:00-20:00 增加外場與內場人力'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 5. AI AGENT ARCHITECTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. AI 人工智慧架構 AI Agent Architecture', level=1)

doc.add_paragraph(
    'Sentinel ECXIP 內建 14 個專業 AI Agent，由 AgentOrchestrator 統一協調，'
    '支援 Sequential 與 Parallel 兩種管線執行模式。'
)

add_colored_table(doc,
    ['AI Agent', '中文名稱', '核心職責'],
    [
        ['RiskAgent', '風險代理', '早期預警偵測、危機升級判斷、風險分數計算'],
        ['VOCAgent', '顧客之聲代理', '情感分析、情緒映射、主題建模、意圖偵測'],
        ['CXAgent', '顧客體驗代理', '旅程摩擦分析、滿意度/NPS/流失風險計算'],
        ['PRAgent', '公關代理', '多渠道公關回應生成 (Google/Threads/PTT/Facebook/新聞稿)'],
        ['LegalAgent', '法務代理', '法律風險評估、合規檢查、證據保存建議'],
        ['KnowledgeAgent', '知識代理', 'SOP/FAQ/培訓材料檢索'],
        ['ExecutiveAgent', '高階主管代理', '高階摘要生成、優先事項排序'],
        ['TrendAgent', '趨勢代理', '趨勢分析、異常偵測、季節性模式識別'],
        ['CompetitorAgent', '競品代理', '競品比較、SWOT 分析、基準報告'],
        ['OperationalAgent', '營運代理', '營運數據關聯、根因假設生成'],
        ['PredictionAgent', '預測代理', '7 日多因子預測、動量計算、信心區間'],
        ['StoreIntelligenceAgent', '門市智慧代理', '門市評分、排行、比較分析'],
        ['LearningAgent', '學習代理', '案例匹配、模式發現、解決方案推薦'],
        ['AICOOAgent', 'AI 營運長代理', '全平台綜合分析、策略建議、晨報生成'],
    ],
    col_widths=[3.5, 3, 9.5]
)

doc.add_paragraph(
    'AI Router 提供 5-tier 模型選擇機制 (Flash → Pro → GPT → Reasoning → Deep Research)，'
    '依據任務複雜度、風險等級、延遲需求與成本，自動選擇最適合的 AI 模型。'
)

doc.add_paragraph(
    '15 個預定義管線 (Pipeline)，包含：危機應變、每日晨報、每週報告、'
    '客戶洞察、風險評估、公關回應、營運分析、預測分析、門市智慧、學習記憶、'
    'COO 簡報等全流程 AI 協作。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 6. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. 技術架構 Technology Stack', level=1)

add_colored_table(doc,
    ['層級', '技術', '版本', '用途'],
    [
        ['前端', 'Vanilla JavaScript + CSS3', '—', 'Apple Frosted Glass SPA，12+ 模組化元件'],
        ['前端字型', 'Plus Jakarta Sans + JetBrains Mono', '—', '現代化企業級字型組合'],
        ['API 框架', 'FastAPI', '0.115+', '高效能非同步 REST API + WebSocket'],
        ['ORM', 'SQLAlchemy 2.0 (async)', '2.0+', '完整 async 支援、型別標註 (Mapped)'],
        ['資料驗證', 'Pydantic v2', '2.0+', '請求/回應資料自動驗證'],
        ['主資料庫', 'PostgreSQL', '16', '31 個資料表、7 個業務領域'],
        ['快取/佇列', 'Redis', '7', '快取、WebSocket Pub/Sub、Celery Broker'],
        ['非同步任務', 'Celery', '5.4', '4 個專用佇列、12 個排程任務'],
        ['認證', 'JWT (python-jose) + bcrypt', '—', 'OAuth2 + 角色權限控制'],
        ['容器化', 'Docker + Docker Compose', '—', '7 個服務容器一鍵部署'],
        ['反向代理', 'Nginx', '1.27', '路由代理、WebSocket 升級、靜態資源快取'],
        ['資料遷移', 'Alembic', '—', 'Async 支援、自動生成遷移腳本'],
        ['國際化', '自建 i18n 框架', '—', '260+ Key，zh-TW + en-US'],
        ['LLM 整合', 'OpenAI / Google Gemini API', '—', 'AI Router 動態選擇與成本優化'],
    ],
    col_widths=[2.5, 4.5, 1.5, 7.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 7. DEPLOYMENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('7. 部署方式 Deployment', level=1)

doc.add_paragraph('支援兩種部署模式：Docker Compose 一鍵部署與本地開發模式。')

doc.add_heading('Docker Compose (推薦)', level=2)
doc.add_paragraph(
    '一個指令即可啟動全部 7 個服務：\n'
    '• postgres (PostgreSQL 16)\n'
    '• redis (Redis 7)\n'
    '• backend (FastAPI on port 8000)\n'
    '• celery_worker (非同步任務處理)\n'
    '• celery_beat (排程任務觸發)\n'
    '• frontend (Nginx static on port 80)\n'
    '• nginx (反向代理 on port 26117)'
)
doc.add_paragraph(
    '啟動指令：\n'
    'cd docker && docker compose up -d\n\n'
    '預覽網址：http://localhost:26117\n'
    'API 文件：http://localhost:8000/docs'
)

doc.add_heading('本地開發', level=2)
doc.add_paragraph(
    'Backend: cd backend && pip install -r requirements.txt && python main.py\n'
    'Frontend: cd frontend && python -m http.server 26117'
)

doc.add_heading('Demo 示範模式', level=2)
doc.add_paragraph(
    '系統預設以 Demo Mode 執行，無需任何 API Key 即可展示全部功能。'
    '所有數據為模擬資料，適合產品展示與客戶 Demo。'
    '設定 DEMO_MODE=false 可切換至正式模式，連接真實 LLM 與爬蟲服務。'
)

doc.add_heading('一鍵產生展示資料', level=2)
doc.add_paragraph(
    'cd backend && python scripts/seed_demo_data.py\n\n'
    '自動產生：1 個測試企業、4 家門市、100 筆 VOC 評論、20 筆危機事件、'
    '30 筆營運數據、10 筆歷史案例、7 天預測資料、晨報與門市排行。'
)

doc.add_heading('Demo 測試帳號', level=2)
doc.add_paragraph(
    'Email: admin@sentinel.ai\n'
    'Password: demo123'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 8. i18n
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. 多語系支援 i18n Internationalization', level=1)

doc.add_paragraph(
    'Sentinel ECXIP 內建完整的 i18n 國際化框架，支援繁體中文 (zh-TW) 與英文 (en-US)。'
)
doc.add_paragraph(
    '• 260+ 組翻譯 Key，覆蓋全部 UI 文字\n'
    '• Header 語言切換器 (繁中 / EN)，即時切換不需刷新頁面\n'
    '• localStorage 保存語系偏好，重整後自動沿用\n'
    '• 三層 Fallback 機制：指定語言 → 英文 → Key 名稱 (不報錯)\n'
    '• HTML data-i18n 屬性系統 (124 處)\n'
    '• 支援新增語言：僅需建立翻譯檔並註冊即可'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 9. SYSTEM HEALTH
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. 系統健康監控 System Health Monitoring', level=1)

doc.add_paragraph(
    '內建完整的系統健康檢查中心，即時監控所有服務狀態：'
)
add_colored_table(doc,
    ['監控項目', 'API 端點', '說明'],
    [
        ['全系統深度檢查', 'GET /api/v1/system/health', '檢查 DB / Redis / Celery Worker / Beat / WebSocket / AI Router / Crawler / Notification'],
        ['快速存活檢查', 'GET /api/v1/system/status', '精簡版 Liveness Probe，回傳 {"status":"ok"}'],
        ['Frontend 儀表板', '側邊欄 Health Widget', '前後端連線狀態即時顯示'],
        ['WebSocket 監控', '/ws/alerts/{user_id}', '即時推送系統警示'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 10. ROADMAP
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. 未來路線圖 Roadmap', level=1)

add_colored_table(doc,
    ['階段', '版本', '狀態', '主要內容'],
    [
        ['Phase 1', 'v1.1', '已完成', 'MVP：品牌監控儀表板、AI 分析沙盒、輿情串流、旅程診斷'],
        ['Phase 2', 'v2.0', '已完成', '企業 SaaS 架構重構：微服務、14 AI Agent、Docker、JWT 認證'],
        ['Phase 3', 'v3.0', '已完成', '智慧平台升級：5 大智慧引擎、5 新 Agent、高階主管決策中心'],
        ['Phase 3.1', 'v3.1', '已完成', 'i18n 國際化：繁體中文/英文介面、語言切換器、260+ Key'],
        ['v1.0.0 MVP', 'v1.0.0', '已完成', 'Enterprise MVP 封版：系統健康檢查、Demo Mode、展示資料產生器'],
        ['Phase 4', 'v4.0', '規劃中', '真實 LLM 整合 (OpenAI/Gemini)、社群爬蟲串接、CI/CD、測試覆蓋'],
        ['Phase 5', 'v5.0', '規劃中', 'Multi-Tenant SaaS、Stripe 金流、白標品牌、自訂 KPI 建構器'],
    ],
    col_widths=[2.5, 1.5, 2, 10]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 11. CLOSING
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. 結語 Closing Statement', level=1)

doc.add_paragraph(
    'Sentinel ECXIP 的核心目標不是建立一套商譽監控工具，而是打造一個真正能夠理解顧客、'
    '協助企業決策、持續學習並創造品牌價值的 Enterprise AI Platform。'
)
doc.add_paragraph(
    '透過 AI Agent、Voice of Customer、Customer Experience、Brand Intelligence、'
    'Operational Intelligence 與 Predictive Analytics，Sentinel ECXIP 將成為企業數位轉型的重要基礎設施，'
    '協助企業從被動回應市場，邁向主動預測、持續優化與智慧經營。'
)
doc.add_paragraph(
    '本平台目前以 Enterprise MVP v1.0.0 版本提供，所有功能模組皆已完成開發與整合測試，'
    '可立即用於產品展示、客戶 Demo 與 AI 競賽展示。'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Sentinel AI ECXIP Enterprise MVP v1.0.0 —')
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x00, 0x71, 0xE3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Ready for Demo & Portfolio')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

# ════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════
output_path = r'D:\Ai study\Aisa_AI\Daily_Ai_002_商譽雷達\Sentinel_ECXIP_專案計劃書_v1.0.0.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
